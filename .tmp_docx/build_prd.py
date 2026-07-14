from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/prd/AI-Hook-Lab-离线评测系统-PRD.md"
OUTPUT = ROOT / "docs/prd/AI-Hook-Lab-离线评测系统-PRD.docx"
ASSET_DIR = ROOT / ".tmp_docx/assets"
SKILL_DIR = Path("/Users/limyoon/.codex/plugins/cache/openai-primary-runtime/documents/26.709.11516/skills/documents")
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402

BODY_FONT = "Calibri"
CJK_FONT = "Arial Unicode MS"
MONO_FONT = "Menlo"
BLACK = RGBColor(28, 31, 36)
MUTED = RGBColor(93, 101, 112)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BRAND_RED = RGBColor(230, 0, 35)  # named override: AI Hook Lab accent
LIGHT_GRAY = "F2F4F7"
LINE_GRAY = "D9DEE6"
CALLOUT_GRAY = "F6F7F9"
WARNING_FILL = "FFF4E5"


def set_run_font(run, name=BODY_FONT, size=11, color=BLACK, bold=None, italic=None, east_asia=CJK_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color="E60023", size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def set_cell_border(cell, color=LINE_GRAY, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_inline(paragraph, text, size=11, color=BLACK):
    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=MONO_FONT, east_asia=CJK_FONT, size=max(9, size - 1), color=DARK_BLUE)
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "EEF1F5")
            run._r.get_or_add_rPr().append(shade)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def table_widths(rows):
    cols = len(rows[0])
    if cols == 2:
        return [2700, 6660]
    if cols == 3:
        return [1800, 3000, 4560]
    lengths = []
    for col in range(cols):
        longest = max(len(re.sub(r"[`*]", "", row[col])) for row in rows)
        lengths.append(min(max(longest, 7), 34))
    return column_widths_from_weights(lengths, 9360)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = table_widths(rows)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if r_idx == 0:
                shade_cell(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if c_idx == 0 and len(row) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value.strip(), size=9.2 if len(rows[0]) >= 4 else 9.6)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = DARK_BLUE
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def load_font(size, bold=False):
    font_path = "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc"
    return ImageFont.truetype(font_path, size=size)


def wrap_text(draw, text, font, max_width):
    lines, current = [], ""
    for ch in text:
        test = current + ch
        if draw.textbbox((0, 0), test, font=font)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_centered(draw, box, text, font, fill, spacing=8):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 36)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + spacing * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + spacing


def arrow(draw, start, end, color="#7B8490", width=5):
    draw.line([start, end], fill=color, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    for delta in (2.55, -2.55):
        point = (end[0] + length * math.cos(angle + delta), end[1] + length * math.sin(angle + delta))
        draw.line([end, point], fill=color, width=width)


def make_main_flow(path):
    img = Image.new("RGB", (1400, 1080), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_font, box_font, owner_font = load_font(42, True), load_font(27, True), load_font(21, True)
    d.text((60, 40), "离线评测全角色主流程", font=title_font, fill="#1C1F24")
    boxes = [
        ("管理员", "初始化账号、案例与 Prompt"),
        ("管理员", "创建批次并锁定快照"),
        ("系统", "生成 120 个任务 / 360 条候选"),
        ("管理员", "筛选 120 条正式结果"),
        ("评测者 A / B", "独立评分与 A/B 盲评"),
        ("系统", "聚合评分并识别分歧"),
        ("裁决者", "匿名处理意向和 A/B 分歧"),
        ("管理员", "Bad Case 根因与动作复盘"),
        ("系统", "计算七项升级门槛"),
        ("管理员", "导出报告并人工决策"),
    ]
    colors = {"管理员": (255, 237, 241), "系统": (235, 242, 249), "评测者 A / B": (238, 247, 242), "裁决者": (255, 246, 224)}
    positions = []
    box_w, box_h, gap_x = 238, 165, 28
    for row in range(2):
        y = 170 + row * 395
        for col in range(5):
            x = 55 + col * (box_w + gap_x)
            positions.append((x, y, x + box_w, y + box_h))
    for idx, ((owner, text), box) in enumerate(zip(boxes, positions)):
        d.rounded_rectangle(box, radius=18, fill=colors[owner], outline="#CBD2DB", width=3)
        x1, y1, x2, _ = box
        d.rounded_rectangle((x1 + 18, y1 + 14, x2 - 18, y1 + 50), radius=14, fill="#FFFFFF")
        ow = d.textbbox((0, 0), owner, font=owner_font)[2]
        d.text((x1 + (box_w - ow) / 2, y1 + 20), owner, font=owner_font, fill="#5C6470")
        draw_centered(d, (x1 + 14, y1 + 56, x2 - 14, y1 + box_h - 12), text, box_font, "#1C1F24")
        if idx < 4:
            arrow(d, (x2 + 4, (y1 + box[3]) // 2), (positions[idx + 1][0] - 4, (positions[idx + 1][1] + positions[idx + 1][3]) // 2))
        elif idx == 4:
            arrow(d, ((x1 + x2) // 2, box[3] + 4), ((positions[5][0] + positions[5][2]) // 2, positions[5][1] - 4))
        elif idx < 9:
            arrow(d, (x2 + 4, (y1 + box[3]) // 2), (positions[idx + 1][0] - 4, (positions[idx + 1][1] + positions[idx + 1][3]) // 2))
    note_font = load_font(23)
    d.rounded_rectangle((55, 950, 1345, 1030), radius=16, fill="#FFF4E5", outline="#E3B258", width=2)
    d.text((82, 975), "证据边界：Mock 只能验证流程，完整 Live + 真人双评 + 60 案例才允许形成升级建议。", font=note_font, fill="#7A5300")
    img.save(path, quality=95)


def make_state_flow(path):
    img = Image.new("RGB", (1400, 780), "#FFFFFF")
    d = ImageDraw.Draw(img)
    title_font, box_font, small_font = load_font(42, True), load_font(27, True), load_font(21)
    d.text((60, 40), "评测批次状态与恢复路径", font=title_font, fill="#1C1F24")
    states = ["draft", "generating", "generated", "selecting", "reviewing", "adjudicating", "completed"]
    boxes = []
    box_w, box_h = 165, 86
    for idx, state in enumerate(states):
        x = 55 + idx * 190
        y = 250 if idx < 5 else 455
        if idx >= 5:
            x = 815 + (idx - 5) * 270
        box = (x, y, x + box_w, y + box_h)
        boxes.append(box)
        fill = "#FFE9EE" if state == "completed" else "#EEF3F8"
        outline = "#E60023" if state == "completed" else "#9EABB9"
        d.rounded_rectangle(box, radius=17, fill=fill, outline=outline, width=3)
        draw_centered(d, box, state, box_font, "#1C1F24")
    labels = ["开始生成", "任务完成", "开始筛选", "正式结果齐全"]
    for i in range(4):
        arrow(d, (boxes[i][2] + 4, (boxes[i][1] + boxes[i][3]) // 2), (boxes[i + 1][0] - 4, (boxes[i + 1][1] + boxes[i + 1][3]) // 2))
        d.text((boxes[i][2] + 18, boxes[i][1] - 34), labels[i], font=small_font, fill="#66707C")
    arrow(d, ((boxes[4][0] + boxes[4][2]) // 2, boxes[4][3] + 4), ((boxes[5][0] + boxes[5][2]) // 2, boxes[5][1] - 4))
    d.text((880, 380), "存在分歧", font=small_font, fill="#66707C")
    arrow(d, (boxes[5][2] + 4, (boxes[5][1] + boxes[5][3]) // 2), (boxes[6][0] - 4, (boxes[6][1] + boxes[6][3]) // 2))
    d.text((1010, 420), "裁决完成", font=small_font, fill="#66707C")
    failed_box = (390, 535, 555, 621)
    d.rounded_rectangle(failed_box, radius=17, fill="#FFF4E5", outline="#D29A31", width=3)
    draw_centered(d, failed_box, "failed", box_font, "#7A5300")
    arrow(d, ((boxes[1][0] + boxes[1][2]) // 2, boxes[1][3] + 4), ((failed_box[0] + failed_box[2]) // 2, failed_box[1] - 4))
    arrow(d, (failed_box[2] + 4, (failed_box[1] + failed_box[3]) // 2), (boxes[1][2] - 20, boxes[1][3] + 14))
    d.text((170, 470), "重试耗尽", font=small_font, fill="#7A5300")
    d.text((565, 575), "管理员单项重试", font=small_font, fill="#7A5300")
    d.rounded_rectangle((55, 680, 1345, 750), radius=16, fill="#F6F7F9", outline="#CBD2DB", width=2)
    d.text((80, 703), "守门规则：终态生成错误、评审未完成或裁决未完成时，批次不得进入正式升级判定。", font=small_font, fill="#4A525D")
    img.save(path, quality=95)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("AI HOOK LAB · PRODUCT REQUIREMENTS")
    set_run_font(run, size=9.5, color=BRAND_RED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("AI Hook Lab 离线评测系统 PRD")
    set_run_font(run, size=25, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("固定案例 · 双人盲评 · 第三人裁决 · Prompt 升级决策")
    set_run_font(run, size=13, color=MUTED)

    metadata = [
        ("文档版本", "V1.0"),
        ("产品状态", "已实现，可进入内部验收"),
        ("适用范围", "Prompt 离线评测、版本升级决策与 Bad Case 复盘"),
        ("更新日期", "2026 年 7 月 13 日"),
        ("文档负责人", "产品 / AI 应用团队"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}：")
        set_run_font(r, size=10.5, color=BLACK, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(14)
    add_bottom_border(rule)

    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(14)
    shade_paragraph(callout, WARNING_FILL)
    r = callout.add_run("核心决策边界  ")
    set_run_font(r, size=10.5, color=RGBColor(122, 83, 0), bold=True)
    r = callout.add_run("系统只提供升级建议；Mock、子集和未完成批次永远不得输出正式升级结论。")
    set_run_font(r, size=10.5, color=RGBColor(122, 83, 0))


def build_docx():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    main_flow = ASSET_DIR / "main-flow.png"
    state_flow = ASSET_DIR / "state-flow.png"
    make_main_flow(main_flow)
    make_state_flow(state_flow)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("AI HOOK LAB · 离线评测系统 PRD")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    add_page_field(footer.paragraphs[0])

    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    diagram_index = 0
    frontmatter_quote_skipped = False
    in_code = False
    code_language = ""
    code_lines = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                text = lines[index].strip()[1:].strip()
                if text:
                    quote_lines.append(text)
                index += 1
            if not frontmatter_quote_skipped:
                frontmatter_quote_skipped = True
                continue
            if quote_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.14)
                p.paragraph_format.right_indent = Inches(0.14)
                p.paragraph_format.space_after = Pt(10)
                shade_paragraph(p, CALLOUT_GRAY)
                add_inline(p, "\n".join(quote_lines), size=9.5, color=MUTED)
            continue
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_language = stripped[3:].strip()
                code_lines = []
            else:
                if code_language == "mermaid":
                    image_path = main_flow if diagram_index == 0 else state_flow
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    picture_width = 5.75 if diagram_index == 0 else 6.0
                    p.add_run().add_picture(str(image_path), width=Inches(picture_width))
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(9)
                    text = "图 1：离线评测全角色主流程" if diagram_index == 0 else "图 2：评测批次状态与恢复路径"
                    r = caption.add_run(text)
                    set_run_font(r, size=9, color=MUTED, italic=True)
                    diagram_index += 1
                elif code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.12)
                    p.paragraph_format.right_indent = Inches(0.12)
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = 1.0
                    shade_paragraph(p, "F3F5F7")
                    r = p.add_run("\n".join(code_lines))
                    set_run_font(r, name=MONO_FONT, east_asia=CJK_FONT, size=8.5, color=RGBColor(40, 48, 58))
                in_code = False
                code_language = ""
                code_lines = []
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("| ") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1].strip()):
            table_lines = [line]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in table_lines]
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            text = heading.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, text, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
            for run in p.runs:
                run.bold = True
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            add_inline(p, stripped)
            index += 1
            continue
        if stripped.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[6:])
            index += 1
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            index += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, stripped)
        index += 1

    props = doc.core_properties
    props.title = "AI Hook Lab 离线评测系统 PRD"
    props.subject = "Prompt 离线评测、版本升级决策与 Bad Case 复盘"
    props.author = "AI Hook Lab 产品团队"
    props.keywords = "AI Hook Lab, PRD, Prompt Evaluation, Offline Evaluation"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
